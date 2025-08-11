# annotator.py
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from typing import List, Dict
import io
import re

HIGHLIGHT_COLOR = RGBColor(255, 230, 150)  # subtle highlight (this sets run font color; python-docx doesn't set highlight easily)

def insert_inline_comment_in_docx(doc_bytes: bytes, issues: List[Dict], original_text: str) -> bytes:
    """
    Creates a copy of the docx where each paragraph containing a detected issue gets inline bracketed comment appended
    and the comment text is highlighted / colored.
    Returns bytes of edited docx.
    """
    doc = Document(io.BytesIO(doc_bytes))
    lower_text = original_text.lower()
    # For each issue we'll search for a snippet in the doc and append a bracketed comment.
    for issue in issues:
        details = issue.get("details", "")
        suggestion = issue.get("recommendation", "")
        snippet = None
        # pick a short anchor from details or issue name
        if details:
            # choose a word from details as anchor if present in doc
            tokens = re.findall(r"[A-Za-z0-9]+", details)
            for t in tokens:
                if t.lower() in lower_text:
                    snippet = t
                    break
        # fallback to issue name
        if not snippet:
            snippet = issue.get("issue", "").split()[0]
        comment_text = f"[[COMMENT - {issue.get('severity','Info')}] {issue.get('issue')}: {suggestion}]"
        # now insert into first paragraph that contains snippet
        inserted = False
        for p in doc.paragraphs:
            if snippet.lower() in p.text.lower():
                p.add_run(" ")
                run = p.add_run(comment_text)
                # color it red-ish to stand out
                run.font.color.rgb = RGBColor(192, 0, 0)
                run.bold = True
                inserted = True
                break
        if not inserted:
            # append to document end
            p = doc.add_paragraph()
            run = p.add_run(comment_text)
            run.font.color.rgb = RGBColor(192, 0, 0)
            run.bold = True
    # return bytes
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

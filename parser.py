# parser.py
import io
from docx import Document
import re
from typing import Dict, List, Tuple

DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "aoa", "articles"],
    "Memorandum of Association": ["memorandum of association", "moa", "memorandum"],
    "Board Resolution": ["board resolution", "resolution of the board"],
    "Shareholder Resolution": ["shareholder resolution", "shareholder's resolution"],
    "Incorporation Application": ["incorporation application", "application for incorporation"],
    "UBO Declaration": ["ubo", "ultimate beneficial owner", "beneficial owner"],
    "Register of Members and Directors": ["register of members", "register of directors", "register of members and directors"],
    "Employment Contract": ["employment contract", "standard employment contract"],
}

def extract_text_from_docx_bytes(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)
    # include table text
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)

def detect_doc_type(text: str) -> Tuple[str, float]:
    txt = text.lower()
    scores = {}
    for dtype, keys in DOC_TYPE_KEYWORDS.items():
        score = sum(txt.count(k) for k in keys)
        scores[dtype] = score
    best = max(scores.items(), key=lambda x: x[1])
    dtype, score = best
    # if low score, return Unknown
    if score == 0:
        return ("Unknown", 0.0)
    # normalized confidence
    conf = score / max(1, sum(scores.values()))
    return (dtype, conf)

def parse_uploaded_docx(file_bytes: bytes) -> Dict:
    text = extract_text_from_docx_bytes(file_bytes)
    doc_type, conf = detect_doc_type(text)
    return {"text": text, "doc_type": doc_type, "doc_confidence": conf}
